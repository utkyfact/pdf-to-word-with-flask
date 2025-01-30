from flask import Flask, render_template, request, send_file
import pdfplumber
from docx import Document
import os
import pytesseract
from PIL import Image
import io
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit
app.config['ALLOWED_EXTENSIONS'] = {'pdf'}

# Tesseract OCR yolu (Sisteminize göre güncelleyin)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows örneği

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def pdf_to_word(pdf_path, docx_path):
    doc = Document()
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Metin çıkarımı
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            else:
                # OCR ile görüntü işleme
                img = page.to_image(resolution=200).original
                img = Image.open(io.BytesIO(img))
                
                # Türkçe ve İngilizce dil desteği
                ocr_text = pytesseract.image_to_string(img, lang='tur+eng')
                doc.add_paragraph(ocr_text)
                
    doc.save(docx_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return render_template('error.html', message='Dosya seçilmedi')
    
    file = request.files['file']
    
    if file.filename == '':
        return render_template('error.html', message='Dosya seçilmedi')
    
    if not allowed_file(file.filename):
        return render_template('error.html', message='Sadece PDF dosyaları yüklenebilir')
    
    try:
        # Dosya işlemleri
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(pdf_path)

        # Dönüştürme
        docx_filename = f"{os.path.splitext(filename)[0]}_converted.docx"
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        
        pdf_to_word(pdf_path, docx_path)

        # Temizlik
        os.remove(pdf_path)
        
        # Dosyayı gönderdikten SONRA silmek için
        response = send_file(
            docx_path,
            as_attachment=True,
            download_name=docx_filename
        )

        # İstek tamamlandıktan sonra temizlik yap
        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
            except Exception as e:
                app.logger.error(f"Dosya silinemedi: {str(e)}")
        
        return response

    except Exception as e:
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)
        return render_template('error.html', message=str(e))
    
if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)